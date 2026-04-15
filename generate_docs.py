from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_architecture_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('EpiChat: System Architecture & Technical Overview', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document outlines the end-to-end architecture of the EpiChat AI-powered epilepsy detection web application.\n')
    
    # 1. High-Level Architecture
    doc.add_heading('1. High-Level Architecture', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('User Flow:\n').bold = True
    p1.add_run('1. User (Patient/Doctor) authenticates & interacts with the React.js Frontend.\n')
    p1.add_run('2. The Frontend communicates with the FastAPI Backend via REST APIs.\n')
    p1.add_run('3. The Backend interacts with the SQLite/PostgreSQL Database for states and user info.\n')
    p1.add_run('4. The Backend handles continuous EEG data by passing it to the ML Pipeline (PyTorch/MNE).\n')
    p1.add_run('5. External APIs (OpenAI for Chatbot, Google Maps for Doctors) are routed securely through the backend.\n')
    
    # 2. Frontend Architecture
    doc.add_heading('2. Frontend Architecture (React.js + Vite)', level=1)
    doc.add_paragraph('The frontend uses a modular component-based architecture:')
    
    ui_list = [
        ('Login Component:', ' Handles user authentication.'),
        ('Dashboard:', ' Central hub containing navigation and global states.'),
        ('Category 1 (Patient History):', ' Displays user medical records and info.'),
        ('Category 2 (EEG Detection):', ' The core module. Handles .edf file uploads, shows live processing state, and renders the Brain Heatmap and Risk Timeline visualizations.'),
        ('Category 3 (Doctors):', ' Fetches nearby neurological specialists based on user location via Google Maps deep links.'),
        ('AI Chatbot Toolbar:', ' A floating widget for real-time medical queries managed via the backend proxy.')
    ]
    for bold_text, normal_text in ui_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
        
    # 3. Backend Architecture
    doc.add_heading('3. Backend Architecture (FastAPI + Python)', level=1)
    doc.add_paragraph('A high-performance Python server chosen for native PyTorch/AI integration.')
    
    api_list = [
        ('/api/login:', ' Authentication middleware & JWT tokens.'),
        ('/api/patient:', ' Profile data retrieval/updates.'),
        ('/api/upload:', ' Core EEG pipeline. Validates .edf files, saves them to temp storage, and triggers the Python inference engine.'),
        ('/api/chat:', ' Receives chat messages, prepends system medical disclaimers, and sends to OpenAI API securely.')
    ]
    for bold_text, normal_text in api_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
        
    # 4. EEG Processing Pipeline
    doc.add_heading('4. EEG Processing Pipeline (The Core Engine)', level=1)
    doc.add_paragraph('Step-by-step flow of how an uploaded file becomes a visual heatmap:')
    
    steps = [
        '1. EDF Upload: User uploads raw continuous EEG data via Category 2.',
        '2. Preprocessing (MNE-Python): Backend applies Bandpass Filtering (1-40Hz), removes artifacts, and segments data into temporal epochs.',
        '3. Feature Extraction: Passes through BIOT Transformer (for global context) and EEGNet 1D-CNN (for local spatial features).',
        '4. Classification: The hybrid model translates tensors into seizure probability scores over time.',
        '5. Visualization Data: The backend returns a JSON payload of Risk % timestamps which React parses into Recharts/Heatmaps.'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')
        
    # 5. Technology Stack
    doc.add_heading('5. Technology Stack', level=1)
    tech = [
        ('Frontend:', ' React.js 18, Vite, Recharts, LocalStorage, CSS Flexbox/Grid.'),
        ('Backend:', ' Python 3.10+, FastAPI, Uvicorn, Pydantic.'),
        ('Machine Learning:', ' PyTorch, MNE-Python, SciPy, NumPy. (Models: BIOT + EEGNet).'),
        ('External APIs:', ' OpenAI API (GPT), Google Maps Deep Links.')
    ]
    for bold_text, normal_text in tech:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
        
    # Save document
    file_path = os.path.join(os.getcwd(), 'EpiChat_Architecture.docx')
    doc.save(file_path)
    print(f"Document generated successfully at {file_path}")

if __name__ == '__main__':
    create_architecture_doc()
